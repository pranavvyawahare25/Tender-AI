"""
Generate realistic sample tender + bidder PDF documents for the CRPF demo.

These are *mock* documents written from scratch — they imitate the structure
and formal tone of real CRPF / GeM tenders for hackathon demo purposes only.
Tender numbers and contents are fictional.
"""
import os
import sys

# Add parent directory to path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
)

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "generated")


# ─── Realistic CRPF tenders for the demo ──────────────────────────────────
TENDERS = [
    {
        "filename":  "CRPF_IT_2025-26_001_Network_Infrastructure.pdf",
        "tender_no": "CRPF/IT/2025-26/001",
        "title":     "Supply, Installation & Commissioning of IT Equipment and Network Infrastructure at CRPF Camps",
        "subject":   ("Supply, installation, and commissioning of IT equipment "
                      "(servers, networking equipment, workstations, UPS) at "
                      "various CRPF establishments across India."),
        "criteria": [
            ("Minimum Annual Turnover",
             "Average annual turnover of Rs. 5 crore in each of the last 3 financial years (2022-23, 2023-24, 2024-25)",
             "CA Certificate / Audited Balance Sheet"),
            ("Similar Projects",
             "Must have completed at least 3 similar projects of IT equipment supply to Central/State Govt organizations in last 5 years",
             "Completion Certificates from clients"),
            ("GST Registration",
             "Must have valid GST registration certificate",
             "Copy of GST Registration Certificate"),
            ("PAN Card",
             "Must have valid PAN card",
             "Copy of PAN Card"),
            ("ISO Certification",
             "Must be ISO 9001:2015 certified",
             "Copy of valid ISO 9001 Certificate"),
            ("EPFO Registration",
             "Must have valid EPFO registration",
             "Copy of EPFO Registration Certificate"),
        ],
        "emd":        "Rs. 5 lakh",
        "experience": "5 years",
    },
    {
        "filename":  "CRPF_MT_2025-26_004_Bullet_Proof_Vehicles.pdf",
        "tender_no": "CRPF/MT/2025-26/004",
        "title":     "Procurement of Bullet Proof Vehicles (BPV) — Phase III for CRPF Operational Battalions",
        "subject":   ("Procurement of 120 Bullet Proof Vehicles (BPV) of B6+ "
                      "ballistic protection level for CRPF operational battalions "
                      "deployed in LWE-affected and J&K sectors."),
        "criteria": [
            ("Minimum Annual Turnover",
             "Average annual turnover of Rs. 100 crore in each of the last 3 financial years",
             "CA Certificate / Audited Balance Sheet"),
            ("OEM Authorization",
             "Bidder must be OEM or hold a valid OEM Authorization Letter for the offered vehicle platform",
             "OEM Authorization Letter"),
            ("Similar Supply Experience",
             "Must have supplied at least 50 BPVs to Central Armed Police Forces (CAPF) or State Police in last 5 years",
             "Purchase Order copies + Performance Certificates"),
            ("BIS / NIJ Test Certification",
             "Vehicle must hold valid B6+/NIJ Level III ballistic test certificate from accredited lab",
             "Test Certificate copy"),
            ("GST Registration",
             "Must have valid GST registration certificate",
             "Copy of GST Registration Certificate"),
            ("Net Worth",
             "Positive net worth in each of the last 3 financial years",
             "CA-certified Net Worth Statement"),
            ("Service Network",
             "Must maintain after-sales service centres in at least 5 states",
             "Service Network Declaration"),
        ],
        "emd":        "Rs. 50 lakh",
        "experience": "10 years",
    },
    {
        "filename":  "CRPF_COMM_2025-26_012_Tactical_Comms.pdf",
        "tender_no": "CRPF/COMM/2025-26/012",
        "title":     "Supply of Tactical Encrypted Communication Sets for Operational Units",
        "subject":   ("Supply of 800 hand-held tactical communication sets with "
                      "AES-256 voice encryption, GPS tracking, and rugged IP-68 "
                      "form factor for CRPF operational units."),
        "criteria": [
            ("Minimum Annual Turnover",
             "Average annual turnover of Rs. 25 crore in each of the last 3 financial years",
             "CA Certificate / Audited Balance Sheet"),
            ("Make-in-India Compliance",
             "Bidder must qualify under Class-I or Class-II Local Supplier as per PPP-MII Order 2017",
             "Self-declaration of local content with CA certification"),
            ("Encryption Certification",
             "Equipment must hold STQC / CCA approval for AES-256 voice encryption",
             "STQC Certificate copy"),
            ("Past Supplies to CAPF",
             "Must have supplied tactical comms to at least 2 CAPF organizations in last 5 years",
             "Purchase Orders + Completion Certificates"),
            ("GST Registration",
             "Must have valid GST registration certificate",
             "Copy of GST Registration Certificate"),
            ("ISO 9001 Certification",
             "Must be ISO 9001:2015 certified",
             "Copy of ISO 9001 Certificate"),
        ],
        "emd":        "Rs. 20 lakh",
        "experience": "7 years",
    },
    {
        "filename":  "CRPF_CIV_2025-26_003_Barracks_Hyderabad.pdf",
        "tender_no": "CRPF/CIV/2025-26/003",
        "title":     "Construction of New Barracks Block at CRPF Group Centre, Hyderabad",
        "subject":   ("Construction of a 4-storey, 240-personnel capacity "
                      "barracks block at CRPF Group Centre, Hyderabad — "
                      "approx. 6,500 sqm built-up area."),
        "criteria": [
            ("Minimum Annual Turnover",
             "Average annual turnover of Rs. 30 crore in each of the last 3 financial years",
             "CA Certificate / Audited Balance Sheet"),
            ("Similar Construction Experience",
             "Completed at least 2 similar single building projects of value Rs. 25 crore each in last 7 years",
             "Completion Certificates from Engineer-in-Charge"),
            ("CPWD/PWD Empanelment",
             "Must be empanelled as Class-I contractor with CPWD or any State PWD",
             "Empanelment Certificate"),
            ("GST Registration",
             "Must have valid GST registration certificate",
             "Copy of GST Registration Certificate"),
            ("PF & ESIC Compliance",
             "Must hold valid EPFO and ESIC registration",
             "Copies of EPFO + ESIC Registration"),
            ("Solvency Certificate",
             "Solvency certificate of Rs. 15 crore from a Scheduled Commercial Bank",
             "Bank Solvency Certificate (not older than 6 months)"),
        ],
        "emd":        "Rs. 30 lakh",
        "experience": "8 years",
    },
    # ─── Karnataka state tender (KTPP Act 1999) ─────────────────────────
    {
        "filename":  "KSP_IT_2025-26_007_State_Police_Network.pdf",
        "tender_no": "KSP/IT/2025-26/007",
        "issuer":    "Karnataka State Police, Government of Karnataka",
        "issuer_addr": "DGP & IGP Office, Nrupathunga Road, Bengaluru — 560001",
        "title":     "Supply &amp; Commissioning of Network Infrastructure for Karnataka State Police Stations",
        "subject":   ("Supply, installation and commissioning of secure network "
                      "infrastructure (routers, switches, fibre, UPS) at 250 "
                      "police stations across Karnataka, under the Karnataka "
                      "State Police IT Modernisation Plan."),
        "framework_lines": [
            "This tender is issued under the <b>Karnataka Transparency in Public "
            "Procurement Act, 1999 (KTPP Act)</b> and its 2000 Rules. The "
            "e-procurement workflow is hosted on <b>eproc.karnataka.gov.in</b> in "
            "compliance with GoK Order DPAR 2/2007.",
            "All bid evaluation will be carried out in line with KTPP Section 4 "
            "(eligibility, pre-qualification, transparency) and CVC's Manual for "
            "Procurement of Goods 2022.",
        ],
        "criteria": [
            ("Minimum Annual Turnover",
             "Average annual turnover of Rs. 10 crore in each of the last 3 financial years",
             "CA Certificate / Audited Balance Sheet"),
            ("Karnataka GST Registration",
             "Must have valid GST registration in Karnataka (29 series)",
             "GST Registration Certificate (Karnataka)"),
            ("e-Procurement Empanelment",
             "Must be registered on eproc.karnataka.gov.in with active vendor ID",
             "e-Procurement vendor registration screenshot"),
            ("Similar Projects",
             "Must have completed at least 2 similar projects of network rollout for State / Central Government in last 5 years",
             "Completion Certificates"),
            ("MSME Preference (Karnataka)",
             "Karnataka-based MSMEs get a 5% price preference under State MSME policy",
             "Udyam Registration + Karnataka domicile certificate"),
            ("ISO 27001 Certification",
             "Must be ISO 27001:2022 certified for information-security management",
             "Copy of ISO 27001 certificate"),
            ("EMD",
             "EMD of Rs. 8 lakh as DD or BG in favour of DGP, Karnataka State Police",
             "DD / Bank Guarantee"),
        ],
        "emd":        "Rs. 8 lakh",
        "experience": "5 years",
    },
]


def _create_doc(filename, content_fn):
    """Helper to create a PDF document."""
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc = SimpleDocTemplate(filepath, pagesize=A4, topMargin=20*mm, bottomMargin=20*mm,
                            leftMargin=15*mm, rightMargin=15*mm)
    styles = getSampleStyleSheet()
    title = ParagraphStyle('DocTitle', parent=styles['Title'], fontSize=16,
                           textColor=colors.HexColor('#1a237e'), spaceAfter=8)
    heading = ParagraphStyle('DocHeading', parent=styles['Heading2'], fontSize=13,
                             textColor=colors.HexColor('#283593'), spaceBefore=14, spaceAfter=6)
    normal = styles['Normal']
    small = ParagraphStyle('Small', parent=normal, fontSize=9, textColor=colors.grey)

    elements = content_fn(title, heading, normal, small)
    doc.build(elements)
    print(f"  Created: {filepath}")
    return filepath


def generate_tender(spec):
    """Generate a tender document from a spec dict (CRPF or state)."""
    def content(title, heading, normal, small):
        els = []
        issuer = spec.get("issuer", "CENTRAL RESERVE POLICE FORCE (CRPF)")
        addr   = spec.get("issuer_addr", "Directorate General — CGO Complex, New Delhi")
        els.append(Paragraph(issuer, title))
        els.append(Paragraph(addr, normal))
        els.append(Spacer(1, 6))
        els.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor('#1a237e')))
        els.append(Spacer(1, 10))

        els.append(Paragraph("NOTICE INVITING TENDER (NIT)", heading))
        els.append(Paragraph(
            f"Tender No: {spec['tender_no']}<br/>"
            f"Date: 01-04-2025<br/>"
            f"Subject: {spec['title']}", normal))
        els.append(Spacer(1, 8))
        els.append(Paragraph(spec["subject"], normal))
        els.append(Spacer(1, 8))
        for line in spec.get("framework_lines", []) or []:
            els.append(Paragraph(line, normal))
            els.append(Spacer(1, 4))
        els.append(Spacer(1, 8))

        els.append(Paragraph("ELIGIBILITY CRITERIA / PRE-QUALIFICATION REQUIREMENTS", heading))
        els.append(Paragraph(
            "The bidder must meet ALL of the following eligibility criteria. "
            "Supporting documents must be enclosed with the Technical Bid.", normal))
        els.append(Spacer(1, 8))

        rows = [["Sl.", "Criterion", "Requirement", "Document Required"]]
        for i, (name, req, doc_) in enumerate(spec["criteria"], start=1):
            rows.append([str(i), name, req, doc_])

        t = Table(rows, colWidths=[28, 110, 215, 130])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a237e')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
        ]))
        els.append(t)
        els.append(Spacer(1, 12))

        els.append(Paragraph("ADDITIONAL CONDITIONS", heading))
        els.append(Paragraph(
            f"• The firm should not be blacklisted or debarred by any Central/State Government organization.<br/>"
            f"• Earnest Money Deposit (EMD) of {spec['emd']} must be submitted.<br/>"
            f"• The bidder must have minimum {spec['experience']} of experience in the relevant domain.<br/>"
            f"• All documents must be self-attested and notarized.", normal))
        els.append(Spacer(1, 12))

        els.append(Paragraph("SCOPE OF WORK", heading))
        els.append(Paragraph(
            "The selected vendor shall be responsible for end-to-end execution "
            "in accordance with the Bill of Quantities (BoQ), the Technical "
            "Specifications, and the Contract Conditions appended to this NIT.", normal))

        els.append(Spacer(1, 20))
        els.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
        els.append(Paragraph("Authorized Signatory — CRPF Procurement Division", small))
        return els

    return _create_doc(spec["filename"], content)


def generate_bidder_1():
    """TechVision Solutions — should PASS all criteria for the IT tender."""
    def content(title, heading, normal, small):
        els = []
        els.append(Paragraph("TechVision Solutions Pvt. Ltd.", title))
        els.append(Paragraph("BID SUBMISSION — Tender No: CRPF/IT/2025-26/001", heading))
        els.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#1a237e')))
        els.append(Spacer(1, 8))

        els.append(Paragraph("COMPANY PROFILE", heading))
        els.append(Paragraph(
            "Company Name: TechVision Solutions Pvt. Ltd.<br/>"
            "Year of Establishment: 2010<br/>"
            "Registered Office: 45, Nehru Place, New Delhi - 110019<br/>"
            "Nature of Business: IT Infrastructure Solutions", normal))
        els.append(Spacer(1, 10))

        els.append(Paragraph("FINANCIAL DETAILS", heading))
        els.append(Paragraph(
            "We hereby certify the following annual turnover details:<br/><br/>"
            "Annual Turnover (2022-23): Rs. 7.5 Crore<br/>"
            "Annual Turnover (2023-24): Rs. 8.1 Crore<br/>"
            "Annual Turnover (2024-25): Rs. 8.0 Crore<br/><br/>"
            "Average Annual Turnover: Rs. 8.0 Crore<br/>"
            "Net Worth (as on 31.03.2025): Rs. 4.2 Crore", normal))
        els.append(Spacer(1, 10))

        els.append(Paragraph("PAST EXPERIENCE — SIMILAR PROJECTS", heading))
        els.append(Paragraph(
            "We have successfully completed 5 similar projects for Government organizations:", normal))

        proj_data = [
            ["#", "Client", "Project", "Value (Rs. Cr)", "Year"],
            ["1", "BSF HQ", "IT Infra Setup", "3.5", "2023"],
            ["2", "CISF Delhi", "Network Equipment", "2.8", "2022"],
            ["3", "NIC Regional", "Server Supply", "4.1", "2024"],
            ["4", "Army Cantonment", "Workstation Supply", "1.9", "2023"],
            ["5", "ITBP Dehradun", "IT Equipment", "2.2", "2024"],
        ]
        t = Table(proj_data, colWidths=[25, 100, 120, 80, 50])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2e7d32')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ]))
        els.append(t)
        els.append(Spacer(1, 10))

        els.append(Paragraph("COMPLIANCE DOCUMENTS", heading))
        els.append(Paragraph(
            "GST Registration No: 07AABCT1234A1ZA — Valid and Active<br/>"
            "PAN Number: AABCT1234A<br/>"
            "ISO 9001:2015 Certification — Valid till 2027<br/>"
            "EPFO Registration — Registered (Establishment Code: DL/12345)<br/>"
            "We declare that our firm has not been blacklisted by any Government organization.", normal))
        els.append(Spacer(1, 20))
        els.append(Paragraph("Authorized Signatory — TechVision Solutions Pvt. Ltd.", small))
        return els

    return _create_doc("bidder_techvision_solutions.pdf", content)


def generate_bidder_2():
    """Bharat Infra Ltd — should FAIL on turnover, projects, ISO."""
    def content(title, heading, normal, small):
        els = []
        els.append(Paragraph("Bharat Infra Ltd.", title))
        els.append(Paragraph("BID SUBMISSION — Tender No: CRPF/IT/2025-26/001", heading))
        els.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#1a237e')))
        els.append(Spacer(1, 8))

        els.append(Paragraph("COMPANY PROFILE", heading))
        els.append(Paragraph(
            "Company Name: Bharat Infra Ltd.<br/>"
            "Year of Establishment: 2018<br/>"
            "Registered Office: 12, MG Road, Bangalore - 560001", normal))
        els.append(Spacer(1, 10))

        els.append(Paragraph("FINANCIAL DETAILS", heading))
        els.append(Paragraph(
            "Annual Turnover (2022-23): Rs. 2.8 Crore<br/>"
            "Annual Turnover (2023-24): Rs. 3.1 Crore<br/>"
            "Annual Turnover (2024-25): Rs. 3.0 Crore<br/><br/>"
            "Average Annual Turnover: Rs. 3.0 Crore", normal))
        els.append(Spacer(1, 10))

        els.append(Paragraph("PAST EXPERIENCE", heading))
        els.append(Paragraph("We have completed 2 projects in IT equipment supply:", normal))
        proj_data = [
            ["#", "Client", "Project", "Value (Rs. Cr)", "Year"],
            ["1", "Municipal Corp", "PC Supply", "0.8", "2023"],
            ["2", "State PWD", "Network Setup", "1.2", "2024"],
        ]
        t = Table(proj_data, colWidths=[25, 100, 120, 80, 50])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#c62828')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ]))
        els.append(t)
        els.append(Spacer(1, 10))

        els.append(Paragraph("COMPLIANCE DOCUMENTS", heading))
        els.append(Paragraph(
            "GST Registration No: 29AABCB5678D1ZP — Valid<br/>"
            "PAN Number: AABCB5678D<br/>"
            "ISO Certification: Not available at present<br/>"
            "EPFO Registration: Application under process<br/>"
            "We declare that our firm has not been blacklisted.", normal))

        els.append(Spacer(1, 20))
        els.append(Paragraph("Authorized Signatory — Bharat Infra Ltd.", small))
        return els

    return _create_doc("bidder_bharat_infra.pdf", content)


def generate_bidder_3():
    """SecureNet Systems — should PASS most, REVIEW on PAN."""
    def content(title, heading, normal, small):
        els = []
        els.append(Paragraph("SecureNet Systems Pvt. Ltd.", title))
        els.append(Paragraph("BID SUBMISSION — Tender No: CRPF/IT/2025-26/001", heading))
        els.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#1a237e')))
        els.append(Spacer(1, 8))

        els.append(Paragraph("COMPANY PROFILE", heading))
        els.append(Paragraph(
            "Company Name: SecureNet Systems Pvt. Ltd.<br/>"
            "Established: 2012<br/>"
            "Office: 78, Cyber City, Gurugram, Haryana - 122002", normal))
        els.append(Spacer(1, 10))

        els.append(Paragraph("FINANCIAL INFORMATION", heading))
        els.append(Paragraph(
            "Our turnover for the last 3 years is as follows:<br/><br/>"
            "Annual Turnover (2022-23): Rs. 5.2 Crore<br/>"
            "Annual Turnover (2023-24): Rs. 5.8 Crore<br/>"
            "Annual Turnover (2024-25): Rs. 5.5 Crore<br/><br/>"
            "Average Annual Turnover: Rs. 5.5 Crore", normal))
        els.append(Spacer(1, 10))

        els.append(Paragraph("SIMILAR PROJECTS COMPLETED", heading))
        els.append(Paragraph(
            "Number of similar projects completed: 3<br/><br/>"
            "1. Delhi Police HQ — Network Infrastructure — Rs. 2.5 Cr (2023)<br/>"
            "2. DRDO Lab — Server and Storage Supply — Rs. 3.8 Cr (2024)<br/>"
            "3. Coast Guard Regional HQ — IT Equipment — Rs. 1.9 Cr (2024)", normal))
        els.append(Spacer(1, 10))

        els.append(Paragraph("COMPLIANCE INFORMATION", heading))
        els.append(Paragraph(
            "GST Registration: 06AABCS9012H1Z5 — Active<br/>"
            "PAN: Document enclosed separately (see annexure)<br/>"
            "ISO 9001:2015 — Certified, valid till December 2026<br/>"
            "EPFO Registration: Registered, Code HR/67890<br/>"
            "Declaration: We have not been blacklisted or debarred by any Government body.", normal))

        els.append(Spacer(1, 20))
        els.append(Paragraph("Authorized Signatory — SecureNet Systems Pvt. Ltd.", small))
        return els

    return _create_doc("bidder_securenet_systems.pdf", content)


def generate_hindi_tender_docx():
    """
    Hindi-language tender (.docx) — exercises the python-docx ingest path
    and the multilingual LLM extractor. Unicode text is stored as-is in
    the DOCX; rendering happens on the viewer's machine where Devanagari
    fonts are present.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        print("  python-docx not installed — skipping Hindi DOCX")
        return None

    out = os.path.join(OUTPUT_DIR, "MHA_HI_2025-26_002_Riot_Gear_Hindi.docx")
    doc = Document()

    title = doc.add_heading("केंद्रीय रिज़र्व पुलिस बल (CRPF)", level=0)
    doc.add_paragraph("गृह मंत्रालय, भारत सरकार · नई दिल्ली")

    doc.add_heading("निविदा सूचना (NIT)", level=1)
    p = doc.add_paragraph()
    p.add_run("निविदा संख्या: CRPF/HI/2025-26/002\n").bold = True
    p.add_run("दिनांक: 01 अप्रैल 2025\n")
    p.add_run("विषय: सीआरपीएफ कर्मियों के लिए दंगा नियंत्रण उपकरण की आपूर्ति "
              "(हेलमेट, बॉडी प्रोटेक्टर, शील्ड, बैटन)।")

    doc.add_heading("पात्रता मानदंड", level=2)
    doc.add_paragraph(
        "बोलीदाता को निम्नलिखित सभी पात्रता मानदंड पूरे करने होंगे। "
        "प्रत्येक मानदंड के लिए सहायक दस्तावेज तकनीकी बोली के साथ संलग्न करें।"
    )

    criteria = [
        ("न्यूनतम वार्षिक टर्नओवर",
         "पिछले 3 वित्तीय वर्षों में औसत वार्षिक टर्नओवर रुपये 8 करोड़ हो।",
         "सीए प्रमाणपत्र / अंकेक्षित बैलेंस शीट"),
        ("समान परियोजनाएँ",
         "केंद्रीय/राज्य सरकार के संगठनों को कम से कम 3 समान आपूर्तियाँ पिछले 5 वर्षों में।",
         "क्रय आदेश + पूर्णता प्रमाणपत्र"),
        ("BIS प्रमाणन",
         "सभी उत्पाद BIS IS 17051 / IS 17066 के अनुसार प्रमाणित होने चाहिए।",
         "BIS प्रमाणपत्र की प्रति"),
        ("GST पंजीकरण",
         "वैध GST पंजीकरण प्रमाणपत्र होना अनिवार्य।",
         "GST प्रमाणपत्र की प्रति"),
        ("ISO 9001 प्रमाणपत्र",
         "ISO 9001:2015 प्रमाणित निर्माता / आपूर्तिकर्ता।",
         "ISO प्रमाणपत्र"),
        ("मेक इन इंडिया",
         "PPP-MII आदेश 2017 के अंतर्गत Class-I / Class-II स्थानीय आपूर्तिकर्ता पात्र।",
         "स्थानीय सामग्री हेतु स्व-घोषणापत्र"),
    ]
    table = doc.add_table(rows=1 + len(criteria), cols=3)
    table.style = "Table Grid"
    h = table.rows[0].cells
    h[0].text = "मानदंड"; h[1].text = "आवश्यकता"; h[2].text = "दस्तावेज"
    for i, (k, req, docu) in enumerate(criteria, start=1):
        r = table.rows[i].cells
        r[0].text = k; r[1].text = req; r[2].text = docu

    doc.add_heading("अतिरिक्त शर्तें", level=2)
    doc.add_paragraph(
        "• किसी भी केंद्रीय/राज्य सरकार के संगठन द्वारा फर्म ब्लैक-लिस्ट नहीं होनी चाहिए।\n"
        "• बयाना राशि (EMD) रुपये 10 लाख जमा करनी होगी।\n"
        "• न्यूनतम 5 वर्षों का अनुभव अनिवार्य।\n"
        "• सभी दस्तावेज स्व-सत्यापित और notarised हों।"
    )
    doc.add_heading("English summary (for the AI extractor)", level=2)
    doc.add_paragraph(
        "Hindi NIT for supply of riot-control equipment to CRPF. "
        "Mandatory eligibility: minimum turnover Rs. 8 crore (3-year average), "
        "3 similar Government supply orders in last 5 years, BIS IS 17051 / IS 17066 "
        "certification, valid GST registration, ISO 9001:2015, PPP-MII Class-I/II "
        "Local Supplier. EMD Rs. 10 lakh."
    )

    doc.save(out)
    print(f"  Created: {out}")
    return out


if __name__ == "__main__":
    print("Generating realistic CRPF + Karnataka demo documents...")
    for spec in TENDERS:
        generate_tender(spec)
    generate_bidder_1()
    generate_bidder_2()
    generate_bidder_3()
    generate_hindi_tender_docx()
    print(f"\nAll documents saved to: {OUTPUT_DIR}")
    print("Done!")
